import os
import json
import tempfile
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document


# -----------------------------
# Configuration
# -----------------------------

load_dotenv()

api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    st.error("GROQ_API_KEY is missing. Please check your environment variables.")
    st.stop()

client = Groq(api_key=api_key)

model = "openai/gpt-oss-120b"


# -----------------------------
# Pydantic Models
# -----------------------------

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


class MatchResult(BaseModel):
    score: float
    details: dict


# -----------------------------
# Read Resume
# -----------------------------

def read_pdf(file_path):
    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def read_docx(file_path):
    document = Document(file_path)

    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


# -----------------------------
# Parse Job Description
# -----------------------------

def parse_job_description(job_description):

    schema = JobD.model_json_schema()

    system_prompt = f"""
You are an expert HR assistant.

Analyze the following job description and extract structured information.

Return ONLY valid JSON matching this schema:

{schema}

Rules:
- Do not return the schema itself.
- Do not invent information.
- If minimum experience is not mentioned, return null.
- If a list has no information, return an empty list.
"""

    user_prompt = f"""
Analyze this job description:

{job_description}
"""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        response_format={
            "type": "json_object"
        }
    )

    data = json.loads(response.choices[0].message.content)

    return JobD(**data)


# -----------------------------
# Parse Resume
# -----------------------------

def parse_resume(resume_text):

    schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert resume parser.

Extract structured information from the resume.

Return ONLY valid JSON matching this schema:

{schema}

Rules:
1. Do not invent information.
2. If a value is not available, return null.
3. If a list has no information, return an empty list.
4. Include internships inside experiences.
5. Extract skills mentioned across the entire resume.
"""

    user_prompt = f"""
Parse the following resume:

{resume_text}
"""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        response_format={
            "type": "json_object"
        }
    )

    data = json.loads(response.choices[0].message.content)

    return Resume(**data)


# -----------------------------
# Match Resume With Job
# -----------------------------

def final_score(job, resume):

    schema = MatchResult.model_json_schema()

    prompt = f"""
You are an HR recruiter.

Compare the candidate's resume with the job description.

JOB DESCRIPTION:

{job.model_dump_json(indent=2)}

CANDIDATE RESUME:

{resume.model_dump_json(indent=2)}

Return JSON matching this schema:

{schema}

Give:
1. Candidate name
2. Matching skills
3. Missing important skills
4. Whether experience requirement is met
5. Overall match percentage from 0 to 100
6. A short final verdict

Keep the response concise and easy to read.
"""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        response_format={
            "type": "json_object"
        }
    )

    data = json.loads(response.choices[0].message.content)

    return MatchResult(**data)


# -----------------------------
# Streamlit UI
# -----------------------------

st.set_page_config(
    page_title="AI Resume Parser & Job Matcher",
    page_icon="📄",
    layout="wide"
)

st.title("AI Resume Parser & Job Matcher")

st.write(
    "Upload multiple resumes and compare them with a job description."
)


# -----------------------------
# Upload Resumes
# -----------------------------

resume_files = st.file_uploader(
    "Upload Resumes",
    type=["pdf", "docx"],
    accept_multiple_files=True
)


# -----------------------------
# Job Description
# -----------------------------

job_description = st.text_area(
    "Job Description",
    height=250,
    placeholder="Paste the job description here..."
)


# -----------------------------
# Analyze
# -----------------------------

if st.button("Analyze Resumes"):

    if not resume_files:

        st.warning("Please upload at least one resume.")

    elif not job_description.strip():

        st.warning("Please enter a job description.")

    else:

        candidates = []

        with st.spinner("Analyzing resumes..."):

            try:

                # Parse job description only once
                job = parse_job_description(job_description)

                # Process every resume
                for resume_file in resume_files:

                    st.write(
                        f"Processing: {resume_file.name}"
                    )

                    suffix = Path(
                        resume_file.name
                    ).suffix.lower()

                    # Save uploaded file temporarily
                    with tempfile.NamedTemporaryFile(
                        delete=False,
                        suffix=suffix
                    ) as temp_file:

                        temp_file.write(
                            resume_file.getbuffer()
                        )

                        temp_path = Path(
                            temp_file.name
                        )

                    try:

                        # Extract resume text
                        if suffix == ".pdf":

                            resume_text = read_pdf(
                                temp_path
                            )

                        else:

                            resume_text = read_docx(
                                temp_path
                            )

                        # Check extracted text
                        if not resume_text.strip():

                            st.warning(
                                f"Could not extract text from {resume_file.name}"
                            )

                            continue

                        # Parse resume
                        resume = parse_resume(
                            resume_text
                        )

                        # Calculate score
                        result = final_score(
                            job,
                            resume
                        )

                        candidates.append(
                            {
                                "filename": resume_file.name,
                                "resume": resume,
                                "result": result
                            }
                        )

                    finally:

                        # Remove temporary file
                        temp_path.unlink(
                            missing_ok=True
                        )

            except Exception as e:

                st.error(
                    f"An error occurred: {str(e)}"
                )

        # -----------------------------
        # Results
        # -----------------------------

        if candidates:

            # Sort highest score first
            candidates.sort(
                key=lambda x: x["result"].score,
                reverse=True
            )

            st.success(
                f"Successfully analyzed {len(candidates)} resume(s)!"
            )

            st.divider()

            # -----------------------------
            # Candidate Ranking
            # -----------------------------

            st.header("🏆 Candidate Ranking")

            for index, candidate in enumerate(
                candidates,
                start=1
            ):

                resume = candidate["resume"]
                result = candidate["result"]

                candidate_name = (
                    resume.name
                    or candidate["filename"]
                )

                st.subheader(
                    f"{index}. {candidate_name}"
                )

                col1, col2 = st.columns(2)

                with col1:

                    st.metric(
                        "Match Score",
                        f"{result.score}%"
                    )

                with col2:

                    st.write(
                        "**Resume:**"
                    )

                    st.write(
                        candidate["filename"]
                    )

                details = result.details

                # -----------------------------
                # Matching Skills
                # -----------------------------

                st.write(
                    "### Matching Skills"
                )

                matching_skills = details.get(
                    "matching_skills",
                    []
                )

                if matching_skills:

                    st.write(
                        ", ".join(
                            matching_skills
                        )
                    )

                else:

                    st.write(
                        "No matching skills found."
                    )

                # -----------------------------
                # Missing Skills
                # -----------------------------

                st.write(
                    "### Missing Important Skills"
                )

                missing_skills = details.get(
                    "missing_important_skills",
                    []
                )

                if missing_skills:

                    st.write(
                        ", ".join(
                            missing_skills
                        )
                    )

                else:

                    st.write(
                        "No major missing skills found."
                    )

                # -----------------------------
                # Experience
                # -----------------------------

                st.write(
                    "### Experience Requirement"
                )

                experience_met = details.get(
                    "experience_requirement_met",
                    "Not specified"
                )

                st.write(
                    experience_met
                )

                # -----------------------------
                # Final Verdict
                # -----------------------------

                st.write(
                    "### Final Verdict"
                )

                verdict = details.get(
                    "final_verdict",
                    "No verdict available."
                )

                st.write(
                    verdict
                )

                st.divider()

        else:

            st.warning(
                "No resumes could be analyzed."
            )